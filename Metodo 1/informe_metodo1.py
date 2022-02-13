import re

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import nltk
from nltk.stem import PorterStemmer


informeName = "informe.txt"
txt_data = open(informeName,"rt")
txt = txt_data.read()
txt_data.close()

terminos = []
coincidencias = []

phraseList = []



resultadoName = "Resultados//" + informeName
resultado = open(resultadoName,"w+")


# Informe inicial
informeName = "Resultados//" + informeName
reportFile = open(informeName,"rt")
txtReport = reportFile.read()
reportFile.close()

txtReport = txtReport.encode('ascii', 'ignore').decode('ascii')



aux = False
auxTerm = False

maxCoincidencia = 3
historial = []

for linea in txt.split("\n"):
	if "HIST" in linea:
		linea = linea.replace("HIST:", "")
		historial.append(linea.strip())

phrases = txtReport.split(".")

for term in txt.split("______________"):
	terms = []
	maxN = []
	if "Encontrado" in term:

		for t_ in term.split("\n"):
			if "Context:" in t_:
				
				t_ = t_.replace("\t", "")
				t_ = t_.replace("Context:", "")
				t_ = t_.strip()

				print("Contexto = ", t_)

				phraseList.append(phrases[int(t_)])
				print(phraseList)
				break

		
		termSt = term.split("Encontrado")[0].replace("\t", "") + "\n"
		resultado.write(termSt)
		terminoF = termSt.replace("Term:  ", "")
		terminoF = terminoF.replace("\n", "")

		terminoF = terminoF.replace("Context", "")

		tokens = nltk.word_tokenize(terminoF)
		tagged = nltk.pos_tag(tokens)
		if len(terminoF.split(" ")) > 1:

			tokens = nltk.word_tokenize(terminoF)
			tagged = nltk.pos_tag(tokens)

			# Marcando todas las palabras de cada termino
			for tag in tagged:
				if tag[1] == "NN":
					terminos.append(tag[0])


		else:
			terminos.append(terminoF)

		for encontrado in term.split("Encontrado"):
			accN = 0
			for sl in encontrado.split("\n"):

				sl = sl.replace("\t", "")
				coinc = re.search('=(.*)%',sl)



				if sl == "\n" or sl == "":
					continue

				maxInterno = 0
				if coinc != None:
					if accN == 0:
						name = sl
						nameC = coinc.group(1)
						nameCNum = float(nameC)
						accN = 1
					elif accN == 1:
						syno = sl
						synoC = coinc.group(1)
						synoCNum = float(synoC)
						accN = 0

						
						if nameCNum >= synoCNum:
							maxInterno = 1
							maxNum = nameCNum 
						else:
							maxInterno = 3
							maxNum = synoCNum 


						if len(maxN) < maxCoincidencia:
							terms.append((name, nameCNum, syno, synoCNum, maxNum))
							maxN.append(max(nameCNum, synoCNum))
							terms.sort(key = lambda x: x[4], reverse = True)
							maxN = sorted(maxN, reverse = True)

						elif max(terms[0][1], terms[0][3]):
							terms.sort(key = lambda x: x[4], reverse = True)
							maxN = sorted(maxN, reverse = True)


		resultado.write("Sorted: ")
		accN = 1
		coincM = 0
		for t in terms:
			

			accS = "\n\t" + str(accN)
			namesS = "\n\t" + t[0]
			synonS = "\n\t" + t[2] + "\n"
			coincN = namesS.split("Coincidencia:")[1]
			coincS = synonS.split("Coincidencia:")[1]
			namesS = namesS.split("Coincidencia:")[0]
			resultado.write(namesS)
			synonS = synonS.split("Coincidencia:")[0]
			resultado.write(synonS)
			pattern = "=(.*?)%"

			substringN = re.search(pattern, coincN).group(1)
			substringS = re.search(pattern, coincS).group(1)
			
			if float(substringN) > float(substringS):
				coincMax = substringN
			else:
				coincMax = substringS

			

		coincidencias.append(coincMax)
		resultado.write("\nCoincidencia: " + coincMax)
	
	else:
		continue	


resultado.close()

# Resultados ordenados 
resultadoName = "Resultados//" + informeName
sortedResult = open(resultadoName,"rt")
txtSorted = sortedResult.read()
sortedResult.close()



document = Document()

document.add_heading('Highlight report')

txtReport = txtReport.encode('ascii', 'ignore').decode('ascii')


paragraph = document.add_paragraph("Los términos encontrados en la ontología Human Phenotype Ontology son resaltados en amarillo, mientras que las frases que contienen información temporal son subrayadas.")

paragraph = document.add_paragraph("")
porter = PorterStemmer()

terminosStemmed = []
for termS in terminos:
	terminosStemmed.append(porter.stem(termS))

report = txtReport.split(".")

start = True

terminos_amarillos = []
terminos_tabla = []


for phrase in report:

	underlineBool = False
	if not start:
		paragraph.add_run(".")
	start = False
	

	phrase2 = phrase.split(" ")
	for word in phrase2:
		word = re.sub(r'[^a-zA-Z0-9]', '', word)
		wordS = porter.stem(word)
		run = paragraph.add_run(word + " ")
		boolT = False
		boolTs = False

		if word in terminos:

			run.font.highlight_color = WD_COLOR_INDEX.YELLOW
			boolT = True
			
			terminos_amarillos.append(word)


		elif wordS in terminosStemmed:
			run.font.highlight_color = WD_COLOR_INDEX.YELLOW
			boolTs = True

			terminos_amarillos.append(word)

		if boolT or boolTs:
			if word in terminos:
				terminos.remove(word)
			if wordS in terminosStemmed:
				terminosStemmed.remove(wordS)

		boolAUX = False


		if phrase.strip() in historial:
			run.underline = True
		else:
			run.underline = False


document.add_page_break()

# Lista de terminos ordenados
document.add_heading('Sorted HPO finds')
paragraph2 = document.add_paragraph("TERMS son las palabras clave que se han buscado en la ontología, debajo aparece el contexto en el que se encuentran. NAMES y SYNONYMS son los significados que se han encontrado en la ontología. Estos están ordenados de forma que el primer par NAMES-SYNONYMS tiene mayor probabilidad de ser el acertado según el contexto. En la columna RELEVANT ha de indicar si considera relevante el término encontrado o no. ")

tables = []
splittedTerm = txtSorted.split("Term:")


titulos = document.add_table(rows=1, cols=4)
cell = titulos.cell(0, 0)
cell.text = "TERMS"
cell = titulos.cell(0, 1)
cell.text = "NAMES"
cell = titulos.cell(0, 2)
cell.text = "SYNONYMS"
cell = titulos.cell(0, 3)
cell.text = "RELEVANT (YES/NO)"


shading_elm_0 = parse_xml(r'<w:shd {} w:fill="2dcac1"/>'.format(nsdecls('w')))
titulos.rows[0].cells[3]._tc.get_or_add_tcPr().append(shading_elm_0)

titulos.style = 'LightShading-Accent1'


nameIdx = 1
accPhrase = 0
namesAA = []

contextos = []

for idx, st in enumerate(splittedTerm):
	if idx == 0:
		continue
	splittedL = st.split("\n")

	auxIdx = 0
	idxFila = 0
	fila = 0
	columna = 0
	aaaux = 0 ##
	auux = []

	for idx2, sll in enumerate(splittedL):

		auxContext = False
		contextoActual = ""

		
		if "Context:" in sll:
			contextoActual = sll
			contextos.append(contextoActual)
			auxContext = True

	
		if idx2 == 0:
			name = sll.strip()

			tables.append(document.add_table(rows=4, cols=4))
			cell = tables[idx-1].cell(0, 0)
			tables[idx-1].style = 'LightShading-Accent1'

			nameL = name.split(" ")
			nameF = ""

			terminos_tabla.append(name)

			auux.append(sll)

			tokens = nltk.word_tokenize(name)
			tagged = nltk.pos_tag(tokens)

			cell.text = name

			contextCell = tables[idx-1].cell(1, 0)



			shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
			tables[idx-1].rows[3].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)

			shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
			tables[idx-1].rows[3].cells[1]._tc.get_or_add_tcPr().append(shading_elm_1)

			shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
			tables[idx-1].rows[3].cells[2]._tc.get_or_add_tcPr().append(shading_elm_1)

			shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
			tables[idx-1].rows[3].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)

			shading_elm_1 = parse_xml(r'<w:shd {} w:fill="2dcac1"/>'.format(nsdecls('w')))
			tables[idx-1].rows[0].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)

			shading_elm_1 = parse_xml(r'<w:shd {} w:fill="2dcac1"/>'.format(nsdecls('w')))
			tables[idx-1].rows[1].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)

			shading_elm_1 = parse_xml(r'<w:shd {} w:fill="2dcac1"/>'.format(nsdecls('w')))
			tables[idx-1].rows[2].cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)

			shading_elm_1 = parse_xml(r'<w:shd {} w:fill="78cce1"/>'.format(nsdecls('w')))
			tables[idx-1].rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)

			shading_elm_1 = parse_xml(r'<w:shd {} w:fill="78cce1"/>'.format(nsdecls('w')))
			tables[idx-1].rows[1].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)




			tables[idx-1].alignment = WD_TABLE_ALIGNMENT.CENTER
			continue
		
		if auxContext:
			contextoActual = contextoActual.replace("Context:", "")
			contextoActual = contextoActual.strip()
			contextCell.text = "Context: \n" + phrases[int(contextoActual)]


		if "Name:" in sll:
			columna = 1
			cell = tables[idx-1].cell(fila, columna)
			sll = sll.replace("Name:", "")
		elif "Synonym" in sll:
			columna = 2
			cell = tables[idx-1].cell(fila, columna)
			sll = sll.replace("Synonym:", "")
			fila += 1
		else:
			continue

		
		tables[idx-1].alignment = WD_TABLE_ALIGNMENT.LEFT
		cell.text = sll



document.add_page_break()

# Cuestionario
document.add_heading('Cuestionario')
paragraph = document.add_paragraph("Marque la opción más adecuada con una X")

table = document.add_table(rows=6, cols=6)

# Escala Linkert
cell = table.cell(0, 1)
cell.text = "Total desacuerdo"
cell = table.cell(0, 2)
cell.text = "En desacuerdo"
cell = table.cell(0, 3)
cell.text = "Neutral"
cell = table.cell(0, 4)
cell.text = "Algo de acuerdo"
cell = table.cell(0, 5)
cell.text = "Muy de acuerdo"

cell = table.cell(1, 0)
cell.text = 'Considera que los resultados son correctos'
cell = table.cell(2, 0)
cell.text = 'Considera que los resultados son relevantes'
cell = table.cell(3, 0)
cell.text = 'Considera que existe algun fallo grave en los resultados'
cell = table.cell(4, 0)
cell.text = 'Considera que emplear estos resultados agiliza el análisis del informe'
cell = table.cell(5, 0)
cell.text = 'Usaría estos resultados como herramienta para su trabajo'

table.style = 'Table Grid'

paragraph = document.add_paragraph("\n\n\nEscriba su opinión sobre las siguientes cuestiones")

table2 = document.add_table(rows=3, cols=2)
cell = table2.cell(0, 0)
cell.text = '¿Hecha en falta algo en los resultados?'
cell = table2.cell(1, 0)
cell.text = '¿Qué fallos graves considera que hay en los resultados?'
cell = table2.cell(2, 0)
cell.text = '¿Mejoraría algún apartado?'


table2.style = 'Table Grid'
document.add_page_break()



docxName = "Resultados//test//InformeTratado" + numInforme + ".docx"
document.save(docxName)